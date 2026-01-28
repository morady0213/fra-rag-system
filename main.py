"""
FRA RAG System - Main Entry Point

This script provides the main interface for the Egyptian Financial Regulatory
Authority (FRA) RAG system. It handles:
1. Document ingestion and indexing
2. Query processing
3. Retrieval and answer generation

Usage:
    python main.py                    # Interactive mode
    python main.py --ingest           # Force re-indexing of documents
    python main.py --query "سؤالك"    # Single query mode

Requirements:
    - Set XAI_API_KEY environment variable for Grok
    - Place documents in data/sample_docs folder
"""

import argparse
import os
import sys
from pathlib import Path
from typing import Optional, List
from loguru import logger

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.absolute()
sys.path.insert(0, str(PROJECT_ROOT))

from config import (
    SAMPLE_DOCS_DIR,
    RAW_PDFS_DIR,
    XAI_API_KEY,
    DEFAULT_TOP_K,
)

# DOCX support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Arabic RTL display support
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    ARABIC_RTL_AVAILABLE = True
except ImportError:
    ARABIC_RTL_AVAILABLE = False


def format_arabic(text: str) -> str:
    """
    Format Arabic text for proper RTL display with connected letters.
    
    Args:
        text: Input text (may contain Arabic)
        
    Returns:
        Reshaped and reordered text for terminal display
    """
    if not ARABIC_RTL_AVAILABLE:
        return text
    
    try:
        # Reshape Arabic letters to connect properly
        reshaped = arabic_reshaper.reshape(text)
        # Reorder for RTL display
        bidi_text = get_display(reshaped)
        return bidi_text
    except Exception:
        return text

from ingestion.ocr_processor import PDFProcessor
from ingestion.chunking import ArabicTextChunker
from ingestion.arabic_utils import normalize_text
from ingestion.metadata_extractor import MetadataExtractor, create_metadata_extractor
from ingestion.hierarchical_chunker import HierarchicalChunker, create_hierarchical_chunker, ParentDocumentRetriever
from rag_engine.vector_store import VectorStore
from rag_engine.retriever import Retriever
from rag_engine.hybrid_retriever import HybridRetriever, ResponseCache
from rag_engine.query_router import QueryRouter, create_query_router
from rag_engine.react_agent import ReActAgent, create_react_agent
from llm_client.grok_client import GrokClient


# Configure logging
logger.remove()
logger.add(
    sys.stderr,
    format="<green>{time:HH:mm:ss}</green> | <level>{level: <8}</level> | <cyan>{message}</cyan>",
    level="INFO",
)


class FRARAGSystem:
    """
    Main RAG System class for FRA documents.
    
    Orchestrates document ingestion, retrieval, and answer generation.
    """
    
    def __init__(
        self,
        docs_dir: Optional[Path] = None,
        pdfs_dir: Optional[Path] = None,
        use_hybrid: bool = True,
        use_reranker: bool = True,
        use_cache: bool = True,
    ):
        """
        Initialize the RAG system.
        
        Args:
            docs_dir: Directory containing sample documents
            pdfs_dir: Directory containing PDF files
            use_hybrid: Enable hybrid search (vector + BM25)
            use_reranker: Enable cross-encoder reranking
            use_cache: Enable response caching
        """
        self.docs_dir = docs_dir or SAMPLE_DOCS_DIR
        self.pdfs_dir = pdfs_dir or RAW_PDFS_DIR
        
        # Store settings
        self.use_hybrid = use_hybrid
        self.use_reranker = use_reranker
        self.use_cache = use_cache
        
        # Initialize components
        self.vector_store = VectorStore()
        self.retriever = Retriever(vector_store=self.vector_store)
        self.chunker = ArabicTextChunker()
        self.pdf_processor = PDFProcessor()
        
        # Initialize metadata extractor (semantic metadata for filtering)
        self.metadata_extractor = create_metadata_extractor(use_llm=False)
        
        # Initialize hierarchical chunker (parent-document retrieval)
        self.hierarchical_chunker = create_hierarchical_chunker(
            min_chunk_size=100,
            max_chunk_size=500,
            index_level="clause",
        )
        
        # Store for hierarchical documents
        self.hierarchical_docs = {}
        
        # Initialize hybrid retriever (enhanced retrieval)
        self.hybrid_retriever = HybridRetriever(
            vector_store=self.vector_store,
            use_reranker=use_reranker,
            use_cache=use_cache,
        )
        
        # Initialize query router (for complex queries)
        self.query_router = create_query_router(self.hybrid_retriever)
        
        # Initialize response cache
        self.cache = ResponseCache() if use_cache else None
        
        # Initialize LLM client (may fail if no API key)
        self.llm_client = None
        self._init_llm_client()
        
        # Initialize ReAct agent for multi-hop reasoning (after LLM client)
        self.react_agent = None
        if self.llm_client:
            self.react_agent = create_react_agent(
                retriever=self.hybrid_retriever,
                llm_client=self.llm_client,
                max_iterations=5,
                language="ar",
            )
            logger.info("ReAct agent initialized for multi-hop reasoning")
        
        logger.info(f"FRA RAG System initialized (hybrid={use_hybrid}, rerank={use_reranker}, cache={use_cache})")
    
    def _init_llm_client(self):
        """Initialize the LLM client if API key is available."""
        try:
            self.llm_client = GrokClient()
            logger.info("Grok LLM client initialized")
        except ValueError as e:
            logger.warning(f"LLM client not initialized: {e}")
            logger.warning("Set XAI_API_KEY to enable answer generation")
    
    def is_indexed(self) -> bool:
        """Check if documents have been indexed."""
        stats = self.vector_store.get_stats()
        return stats["document_count"] > 0
    
    def ingest_documents(self, force: bool = False) -> int:
        """
        Ingest documents from the sample_docs and raw_pdfs directories.
        
        Args:
            force: If True, reindex even if documents exist
            
        Returns:
            Number of documents indexed
        """
        if self.is_indexed() and not force:
            logger.info(
                f"Vector store already has {self.vector_store.get_stats()['document_count']} documents. "
                "Use --ingest flag to force re-indexing."
            )
            return 0
        
        if force:
            logger.warning("Force re-indexing: clearing existing documents")
            self.vector_store.delete_collection()
        
        all_chunks = []
        
        # Process text/markdown files from sample_docs
        text_chunks = self._process_text_files()
        all_chunks.extend(text_chunks)
        
        # Process DOCX files
        docx_chunks = self._process_docx_files()
        all_chunks.extend(docx_chunks)
        
        # Process PDF files
        pdf_chunks = self._process_pdf_files()
        all_chunks.extend(pdf_chunks)
        
        if not all_chunks:
            logger.warning("No documents found to ingest!")
            logger.info(f"Place documents in: {self.docs_dir}")
            logger.info(f"Or PDF files in: {self.pdfs_dir}")
            return 0
        
        # Add to vector store
        logger.info(f"Indexing {len(all_chunks)} chunks...")
        added = self.vector_store.add_documents(all_chunks)
        
        logger.info(f"Successfully indexed {added} document chunks")
        return added
    
    def _process_text_files(self) -> List[dict]:
        """Process text and markdown files."""
        chunks = []
        
        # Supported text file extensions
        extensions = ["*.txt", "*.md", "*.markdown"]
        
        for ext in extensions:
            for filepath in self.docs_dir.glob(ext):
                try:
                    logger.info(f"Processing: {filepath.name}")
                    
                    with open(filepath, "r", encoding="utf-8") as f:
                        content = f.read()
                    
                    if not content.strip():
                        continue
                    
                    # Chunk the document
                    chunked = self.chunker.chunk_text(
                        text=content,
                        source=filepath.name,
                        metadata={"type": "text", "path": str(filepath)},
                    )
                    
                    for chunk in chunked.chunks:
                        chunks.append({
                            "text": chunk.content,
                            "metadata": chunk.metadata,
                        })
                    
                except Exception as e:
                    logger.error(f"Error processing {filepath}: {e}")
        
        logger.info(f"Processed {len(chunks)} chunks from text files")
        return chunks
    
    def _process_docx_files(self) -> List[dict]:
        """Process DOCX (Word) files."""
        chunks = []
        
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not installed. Skipping DOCX files.")
            return chunks
        
        for filepath in self.docs_dir.glob("*.docx"):
            try:
                logger.info(f"Processing DOCX: {filepath.name}")
                
                # Extract text from DOCX
                doc = DocxDocument(filepath)
                
                # Extract all paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            paragraphs.append(" | ".join(row_text))
                
                content = "\n\n".join(paragraphs)
                
                if not content.strip():
                    logger.warning(f"No text extracted from: {filepath.name}")
                    continue
                
                # Extract semantic metadata
                doc_metadata = self.metadata_extractor.extract(content, filepath.name)
                
                # Chunk the document
                chunked = self.chunker.chunk_text(
                    text=content,
                    source=filepath.name,
                    metadata={
                        "type": "docx",
                        "path": str(filepath),
                        "document_type": doc_metadata.document_type,
                        "entity_types": doc_metadata.entity_types,
                        "topics": doc_metadata.topics,
                        "has_penalties": doc_metadata.has_penalties,
                        "has_capital_requirements": doc_metadata.has_capital_requirements,
                        "has_licensing_requirements": doc_metadata.has_licensing_requirements,
                        "has_branch_requirements": doc_metadata.has_branch_requirements,
                    },
                )
                
                for chunk in chunked.chunks:
                    chunks.append({
                        "text": chunk.content,
                        "metadata": chunk.metadata,
                    })
                
            except Exception as e:
                logger.error(f"Error processing {filepath}: {e}")
        
        logger.info(f"Processed {len(chunks)} chunks from DOCX files")
        return chunks
    
    def _process_pdf_files(self) -> List[dict]:
        """Process PDF files."""
        chunks = []
        
        # Check both directories for PDFs
        pdf_dirs = [self.docs_dir, self.pdfs_dir]
        
        for pdf_dir in pdf_dirs:
            if not pdf_dir.exists():
                continue
            
            for filepath in pdf_dir.glob("*.pdf"):
                try:
                    logger.info(f"Processing PDF: {filepath.name}")
                    
                    # Extract text from PDF
                    doc = self.pdf_processor.process_file(filepath)
                    content = doc.full_text
                    
                    if not content.strip():
                        logger.warning(f"No text extracted from: {filepath.name}")
                        continue
                    
                    # Chunk the document
                    chunked = self.chunker.chunk_text(
                        text=content,
                        source=filepath.name,
                        metadata={
                            "type": "pdf",
                            "path": str(filepath),
                            "pages": doc.total_pages,
                        },
                    )
                    
                    for chunk in chunked.chunks:
                        chunks.append({
                            "text": chunk.content,
                            "metadata": chunk.metadata,
                        })
                    
                except Exception as e:
                    logger.error(f"Error processing {filepath}: {e}")
        
        logger.info(f"Processed {len(chunks)} chunks from PDF files")
        return chunks
    
    def query(
        self,
        question: str,
        k: int = DEFAULT_TOP_K,
        show_sources: bool = True,
        use_hybrid: bool = None,
        use_rerank: bool = None,
    ) -> str:
        """
        Query the RAG system and get an answer.
        
        Args:
            question: User's question (Arabic or English)
            k: Number of documents to retrieve
            show_sources: Whether to display source documents
            use_hybrid: Override hybrid search setting (None = use default)
            use_rerank: Override reranking setting (None = use default)
            
        Returns:
            Generated answer
        """
        if not question.strip():
            return "الرجاء إدخال سؤال."  # Please enter a question
        
        logger.info(f"Query: {question[:50]}...")
        
        # Determine retrieval settings
        hybrid = use_hybrid if use_hybrid is not None else self.use_hybrid
        rerank = use_rerank if use_rerank is not None else self.use_reranker
        
        # Use hybrid retriever if enabled
        if hybrid:
            retrieval_result = self.hybrid_retriever.retrieve_with_context(question, k=k)
        else:
            retrieval_result = self.retriever.retrieve_with_context(question, k=k)
        
        context = retrieval_result["context"]
        sources = retrieval_result["sources"]
        
        if show_sources and sources:
            print("\n" + "=" * 60)
            print(format_arabic("📚 المصادر المسترجعة (Retrieved Sources):"))
            print("-" * 60)
            for i, source in enumerate(sources, 1):
                print(format_arabic(f"  {i}. {source['source']} (relevance: {source['score']:.2%})"))
            print("=" * 60 + "\n")
        
        if not context:
            return "لم يتم العثور على معلومات ذات صلة في الوثائق المتاحة."
        
        # Generate answer
        if self.llm_client:
            try:
                result = self.llm_client.generate(
                    query=question,
                    context=context,
                    sources=[s["source"] for s in sources],
                )
                return result.answer
            except Exception as e:
                logger.error(f"LLM generation failed: {e}")
                return f"خطأ في توليد الإجابة: {e}\n\nالسياق المسترجع:\n{context}"
        else:
            # No LLM client - return raw context
            return f"⚠️ LLM غير متاح (قم بتعيين XAI_API_KEY)\n\nالسياق المسترجع:\n{context}"
    
    def interactive_mode(self):
        """Run interactive query mode."""
        print("\n" + "=" * 60)
        print(format_arabic("🏛️  نظام الاسترجاع المعزز للهيئة العامة للرقابة المالية"))
        print("    FRA RAG System - Interactive Mode")
        print("=" * 60)
        print(format_arabic("\nاكتب سؤالك باللغة العربية أو الإنجليزية."))
        print("Type your question in Arabic or English.")
        print(format_arabic("Type 'exit' or 'خروج' to quit.\n"))
        
        while True:
            try:
                question = input(format_arabic("❓ السؤال: ")).strip()
                
                if not question:
                    continue
                
                if question.lower() in ["exit", "quit", "خروج", "q"]:
                    print(format_arabic("\nشكراً لاستخدام النظام. مع السلامة!"))
                    break
                
                # Get answer
                answer = self.query(question)
                
                print("\n" + "-" * 60)
                print(format_arabic("📝 الإجابة:"))
                print("-" * 60)
                print(format_arabic(answer))
                print("\n")
                
            except KeyboardInterrupt:
                print(format_arabic("\n\nتم إيقاف النظام."))
                break
            except Exception as e:
                logger.error(f"Error: {e}")
                print(format_arabic(f"حدث خطأ: {e}\n"))


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="FRA RAG System - Egyptian Financial Regulatory Authority",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python main.py                              # Interactive mode
    python main.py --ingest                     # Re-index documents
    python main.py --query "ما هي اختصاصات الهيئة؟"  # Single query
    python main.py --stats                      # Show statistics
        """,
    )
    
    parser.add_argument(
        "--ingest",
        action="store_true",
        help="Force re-indexing of all documents",
    )
    
    parser.add_argument(
        "--query", "-q",
        type=str,
        help="Single query mode - ask a question directly",
    )
    
    parser.add_argument(
        "--stats",
        action="store_true",
        help="Show vector store statistics",
    )
    
    parser.add_argument(
        "--k",
        type=int,
        default=DEFAULT_TOP_K,
        help=f"Number of documents to retrieve (default: {DEFAULT_TOP_K})",
    )
    
    args = parser.parse_args()
    
    # Initialize system
    system = FRARAGSystem()
    
    # Handle --stats
    if args.stats:
        stats = system.vector_store.get_stats()
        print("\n📊 Vector Store Statistics:")
        print("-" * 40)
        for key, value in stats.items():
            print(f"  {key}: {value}")
        print()
        return
    
    # Handle --ingest
    if args.ingest:
        system.ingest_documents(force=True)
        return
    
    # Check if we need to ingest first
    if not system.is_indexed():
        print("\n⚠️  لم يتم العثور على وثائق مفهرسة.")
        print("   No indexed documents found.")
        print(f"\n   Place documents in: {SAMPLE_DOCS_DIR}")
        print(f"   Or PDF files in: {RAW_PDFS_DIR}")
        
        # Check if there are files to ingest
        has_docs = any(SAMPLE_DOCS_DIR.glob("*")) if SAMPLE_DOCS_DIR.exists() else False
        has_pdfs = any(RAW_PDFS_DIR.glob("*.pdf")) if RAW_PDFS_DIR.exists() else False
        
        if has_docs or has_pdfs:
            print("\n   Found documents to ingest. Indexing now...")
            system.ingest_documents()
        else:
            print("\n   Add documents and run again, or use --ingest flag.")
            return
    
    # Handle --query
    if args.query:
        answer = system.query(args.query, k=args.k)
        print("\n" + "=" * 60)
        print(format_arabic("📝 الإجابة:"))
        print("=" * 60)
        print(format_arabic(answer))
        print()
        return
    
    # Default: interactive mode
    system.interactive_mode()


if __name__ == "__main__":
    main()
